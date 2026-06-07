import json
import io
import re
from typing import List, Tuple
from datetime import datetime
from urllib.parse import quote

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from sqlmodel import Session, select
from openai import AsyncOpenAI

from app.config import LLM_API_KEY, LLM_BASE_URL, LLM_MODEL
from app.models import Record, RecordCreateInfo, Signature
from app.database import get_session
from app.api.deps import verify_token
from app.schemas.record import RecordStartRequest, ChatRequest, ChatResponse, ChatAudioResponse, SignatureRequest
from app.prompts import get_interrogation_prompt
from app.api.routers.audio import transcribe_upload_file

router = APIRouter(prefix="/records", tags=["AI笔录核心业务"])


STATUS_WITNESS_OPENING_1 = "证人_告知如实作证义务"
STATUS_WITNESS_OPENING_2 = "证人_宣读权利义务告知书"
STATUS_WITNESS_OPENING_3 = "证人_确认健康状况"
STATUS_WITNESS_OPENING_4 = "证人_确认可接受询问"
STATUS_WITNESS_OPENING_5 = "证人_采集个人信息"
STATUS_CASE_STATEMENT = "案情陈述中"
STATUS_AI_ASKING = "AI询问中"
STATUS_MANUAL_INTERVENTION = "人工干预"
STATUS_FINISHED = "笔录结束"

OPENING_STEPS = {
    STATUS_WITNESS_OPENING_1: "我们是新城区公安分局刑警大队的民警（出示人民警察证），现依法向你询问有关问题。根据刑事诉讼法的有关规定，你应当如实提供证据、证言，如果有意作伪证或者隐匿罪证的，要负法律责任。你明白吗？",
    STATUS_WITNESS_OPENING_2: "现在依法向你宣读《证人诉讼权利义务告知书》：你有权使用本民族语言文字进行诉讼；有权核对询问笔录，认为记录有遗漏或者差错的，可以提出补充或者改正；有权对与本案无关的问题拒绝回答；有权要求侦查人员回避。你应当如实提供证据、证言，故意作伪证或者隐匿罪证的，应当承担相应法律责任。以上权利义务你是否听清楚？",
    STATUS_WITNESS_OPENING_3: "你是否患有严重疾病或其他不适宜作证的情况？",
    STATUS_WITNESS_OPENING_4: "现在开始采集你的个人情况。",
    STATUS_WITNESS_OPENING_5: "请陈述你的个人情况：姓名、性别、年龄、出生日期、身份证件号码、是否为人大代表、现住址、联系方式、户籍所在地。"
}
FIXED_CLOSING = "你还有什么需要补充说明的吗？如果以上笔录核对无误，请仔细阅读后签名按手印。"


client = AsyncOpenAI(
    api_key=LLM_API_KEY,
    base_url=LLM_BASE_URL
)


def contains_any(text: str, keywords: List[str]) -> bool:
    normalized_text = text.strip()
    return any(keyword in normalized_text for keyword in keywords)


def contains_none(text: str, keywords: List[str]) -> bool:
    normalized_text = text.strip()
    return not any(keyword in normalized_text for keyword in keywords)


def is_clear_yes(text: str, yes_keywords: List[str], deny_keywords: List[str]) -> bool:
    normalized_text = text.strip()
    return contains_any(normalized_text, yes_keywords) and contains_none(normalized_text, deny_keywords)


def is_clear_no(text: str, no_keywords: List[str], exclude_keywords: List[str]) -> bool:
    normalized_text = text.strip()
    return contains_any(normalized_text, no_keywords) and contains_none(normalized_text, exclude_keywords)


def parse_extracted_info(extracted_info: str) -> dict:
    try:
        return json.loads(extracted_info) if extracted_info else {}
    except json.JSONDecodeError:
        return {}


def build_initial_extracted_info(req: RecordStartRequest) -> dict:
    return {
        "姓名": req.person_name.strip(),
        "性别": "",
        "年龄": "",
        "出生日期": "",
        "身份证件号码": req.id_number.strip() if req.id_type.strip() in ["身份证", "居民身份证"] else "",
        "是否为人大代表": "",
        "现住址": "",
        "联系方式": "",
        "户籍所在地": "",
        "案情": req.case_name.strip(),
        "发生时间": "",
        "发生地点": "",
        "案发经过": "",
        "嫌疑人特征": "",
        "作案工具及涉案物品": "",
        "其他线索": "",
        "相关人员信息": ""
    }


def build_initial_context(req: RecordStartRequest) -> str:
    context_lines = [
        f"案件类型：{req.case_type.strip() or '未填写'}",
        f"案件名称：{req.case_name.strip() or '未填写'}",
        f"被询问人身份：{req.person_type.strip() or '未填写'}",
        f"被询问人姓名：{req.person_name.strip() or '未填写'}",
        f"证件类型：{req.id_type.strip() or '未填写'}",
        f"证件号码：{req.id_number.strip() or '未填写'}"
    ]
    return "案件基础信息：\n" + "\n".join(context_lines)



def get_case_and_person_type(session: Session, record_id: int) -> Tuple[str, str]:
    """获取案件类型和人员身份"""
    statement = select(RecordCreateInfo).where(RecordCreateInfo.record_id == record_id)
    create_info = session.exec(statement).first()
    if create_info:
        case_type = create_info.case_type.strip() if create_info.case_type else "盗窃案"
        person_type = create_info.person_type.strip() if create_info.person_type else "受害人"
        return case_type, person_type
    return "盗窃案", "受害人"


def get_respondent_label(session: Session, record_id: int) -> str:
    statement = select(RecordCreateInfo).where(RecordCreateInfo.record_id == record_id)
    create_info = session.exec(statement).first()
    if create_info and create_info.person_type.strip():
        return create_info.person_type.strip()
    return "被询问人"


def handle_opening_flow(current_status: str, user_text: str) -> Tuple[str, str, bool]:
    # ==========================================
    # 状态 1：告知如实作证义务
    # ==========================================
    if current_status == STATUS_WITNESS_OPENING_1:
        # 1. 乖乖配合，流程推进
        if is_clear_yes(user_text, ["明白", "听明白", "清楚", "知道了", "理解了", "嗯", "对", "行"], ["不"]):
            return OPENING_STEPS[STATUS_WITNESS_OPENING_2], STATUS_WITNESS_OPENING_2, True
        # 2. 明确拒绝/抗拒（非法回答）-> 严厉警告，状态原地锁定！
        elif is_clear_no(user_text, ["不明白", "没明白", "不清楚", "不理解", "不知道", "不想", "不愿"], []):
            return "【严厉提醒】作证是每个公民的法定义务。如果有意作伪证或者隐匿罪证，将面临三年以下有期徒刑等法律处罚。请明确回答，你现在听明白了吗？", current_status, True
        # 3. 顾左右而言他（非法回答）-> 纠正话术，状态原地锁定！
        else:
            return "请你针对我的问题明确回答“明白”或“不明白”。根据法律规定，你应当如实提供证据、证言。你明白了吗？", current_status, True

    # ==========================================
    # 状态 2：宣读权利义务告知书
    # ==========================================
    elif current_status == STATUS_WITNESS_OPENING_2:
        if is_clear_yes(user_text, ["听清楚", "清楚", "明白", "知道了", "理解了", "嗯", "对", "是"], ["不"]):
            return OPENING_STEPS[STATUS_WITNESS_OPENING_3], STATUS_WITNESS_OPENING_3, True
        else:
            return "我再向你宣读一遍《证人诉讼权利义务告知书》：你有权使用本民族语言文字进行诉讼；有权核对询问笔录并提出补充或者改正；有权对与本案无关的问题拒绝回答；有权要求侦查人员回避。你应当如实提供证据、证言，故意作伪证或者隐匿罪证的，应当承担相应法律责任。以上权利义务你是否听清楚？", current_status, True

    # ==========================================
    # 状态 3：确认健康状况
    # ==========================================
    elif current_status == STATUS_WITNESS_OPENING_3:
        if contains_any(user_text, ["能坚持", "可以继续", "继续", "不用", "不需要", "没事", "还能", "可以"]):
            return OPENING_STEPS[STATUS_WITNESS_OPENING_5], STATUS_WITNESS_OPENING_5, True
        # 注意：这里的肯定代表“没病”，推进流程
        if contains_any(user_text, ["没有", "无", "没病", "挺好", "健康", "正常", "没"]):
            return OPENING_STEPS[STATUS_WITNESS_OPENING_5], STATUS_WITNESS_OPENING_5, True
        # 这里的否定代表“有病”，给出医疗选项，并将状态推进到下一个去确认“能否接受询问”
        elif contains_any(user_text, ["有", "头晕", "发烧", "心脏病", "不舒服", "病", "疼", "难受"]):
            return "如果你目前身体极度不适，我们可以为你呼叫120医疗援助并暂停询问。请问你是否需要暂停询问或者医疗帮助？", current_status, True
        else:
            return "请明确说明你是否有严重疾病或其他不适宜作证的情况？（如确无异常，请回答“没有”）", current_status, True

    # ==========================================
    # 状态 4：确认可接受询问
    # ==========================================
    elif current_status == STATUS_WITNESS_OPENING_4:
        return OPENING_STEPS[STATUS_WITNESS_OPENING_5], STATUS_WITNESS_OPENING_5, True

    # ==========================================
    # 状态 5：采集个人信息 (大模型交接点！)
    # ==========================================
    elif current_status == STATUS_WITNESS_OPENING_5:
        # 拦截极其不配合的抗拒态度（非法回答）
        if contains_any(user_text, ["不想说", "不愿", "保密", "不知道", "凭什么", "不告诉你", "隐私"]):
            return "【法制教育】配合调查并如实提供真实身份信息，是公民法定义务。拒绝提供或提供虚假信息将承担相应法律后果。请如实陈述你的姓名、出生日期、住址及联系方式等信息。", current_status, True
        
        # 🌟 魔法交接点：只要他不拒绝，我们就认为他开始报个人信息了。
        # 此时，我们返回 matched = False。
        # 控制权将彻底移交给你的 Qwen 大模型！大模型会读取他的信息，提取出 JSON 字段，并极其丝滑地将状态推进到“AI询问中”。
        else:
            return "", current_status, False

    # 兜底：不在开场状态的，统统交给大模型
    return "", current_status, False


@router.post("/start", summary="1. 开启新笔录 (播报开场固定内容)")
async def start_record(
    req: RecordStartRequest,
    token: str = Depends(verify_token),
    session: Session = Depends(get_session)
):
    parts = token.split("_")
    police_number = parts[2] if len(parts) >= 3 else "unknown_police"

    initial_info = build_initial_extracted_info(req)
    opening_text = OPENING_STEPS[STATUS_WITNESS_OPENING_1]
    initial_context = build_initial_context(req)
    record_title = req.case_name.strip() or "AI智能笔录"

    new_record = Record(
        police_number=police_number,
        title=record_title,
        content=f"{initial_context}\nAI警官：{opening_text}",
        status=STATUS_WITNESS_OPENING_1,
        extracted_info=json.dumps(initial_info, ensure_ascii=False)
    )

    session.add(new_record)
    session.commit()
    session.refresh(new_record)

    create_info = RecordCreateInfo(
        record_id=new_record.id,
        police_number=police_number,
        case_type=req.case_type.strip(),
        case_name=req.case_name.strip(),
        person_type=req.person_type.strip(),
        person_name=req.person_name.strip(),
        id_type=req.id_type.strip(),
        id_number=req.id_number.strip()
    )

    session.add(create_info)
    session.commit()

    return {
        "code": 200,
        "message": "笔录初始化成功",
        "data": {
            "record_id": new_record.id,
            "ai_reply": opening_text,
            "status": new_record.status
        }
    }


async def process_chat_message(record_id: int, reporter_text: str, session: Session) -> ChatResponse:
    db_record = session.get(Record, record_id)
    if not db_record:
        raise HTTPException(status_code=404, detail="找不到该笔录ID")

    user_text = reporter_text.strip()
    if not user_text:
        raise HTTPException(status_code=400, detail="识别内容为空，请重新录音")

    current_status = db_record.status
    extracted_info = db_record.extracted_info
    respondent_label = get_respondent_label(session, record_id)
    case_type, person_type = get_case_and_person_type(session, record_id)

    current_extracted_info = parse_extracted_info(extracted_info)
    ai_reply = ""
    new_status = current_status
    new_extracted_info = current_extracted_info

    # 判断是否为个人信息采集阶段
    is_collecting_person_info = current_status == STATUS_WITNESS_OPENING_5

    rule_reply, rule_status, matched = handle_opening_flow(current_status, user_text)

    if matched:
        ai_reply = rule_reply
        new_status = rule_status
        # 规则匹配时，直接将书面化的语音内容写入笔录
        formalized_text = await formalize_text_async(user_text)
        db_record.content += f"\n{respondent_label}：{formalized_text}"
    else:
        system_prompt = get_interrogation_prompt(
            current_status,
            current_extracted_info,
            case_type=case_type,
            person_type=person_type
        )

        try:
            response = await client.chat.completions.create(
                model=LLM_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"以下是完整的聊天记录，请分析并给出下一步回应：\n{db_record.content}\n当前被询问人回答：{user_text}"}
                ],
                temperature=0.1,
                response_format={"type": "json_object"},
                max_tokens=800
            )

            ai_response_text = response.choices[0].message.content
            print(f"🕵️‍♂️ 大模型原始回复内容是：\n{ai_response_text}\n")

            start_idx = ai_response_text.find('{')
            end_idx = ai_response_text.rfind('}')

            if start_idx != -1 and end_idx != -1:
                clean_json_str = ai_response_text[start_idx:end_idx+1]
                result_data = json.loads(clean_json_str)
            else:
                print("❌ 警告：大模型没有返回有效的 JSON 结构！")
                result_data = {
                    "ai_reply": ai_response_text,
                    "new_status": current_status,
                    "extracted_info": current_extracted_info
                }

            ai_reply = result_data.get("ai_reply") or result_data.get("reply") or result_data.get("text", "（大模型正在思考案件，请稍候）")
            new_status = result_data.get("new_status") or result_data.get("status", current_status)
            new_extracted_info = result_data.get("extracted_info") or current_extracted_info

        except Exception as e:
            print(f"大模型调用失败: {e}")
            raise HTTPException(status_code=500, detail="AI 大脑暂时开小差了，请稍后再试。")

        # 根据阶段决定如何写入被询问人的回答
        if is_collecting_person_info:
            # 个人信息采集阶段：用提取出的信息生成书面化个人介绍
            person_intro = generate_person_intro(new_extracted_info, person_type)
            if person_intro:
                db_record.content += f"\n{respondent_label}：{person_intro}"
            else:
                # 如果还没提取到足够信息，先用书面化原文
                formalized_text = await formalize_text_async(user_text)
                db_record.content += f"\n{respondent_label}：{formalized_text}"
        else:
            # 其他阶段：直接将语音转文字内容书面化后写入笔录
            formalized_text = await formalize_text_async(user_text)
            db_record.content += f"\n{respondent_label}：{formalized_text}"

    db_record.content += f"\nAI警官：{ai_reply}"
    db_record.status = new_status
    db_record.extracted_info = json.dumps(new_extracted_info, ensure_ascii=False)

    session.add(db_record)
    session.commit()

    return ChatResponse(
        ai_reply=ai_reply,
        status=db_record.status,
        extracted_info=new_extracted_info
    )


@router.post("/chat", response_model=ChatResponse, summary="2. 核心对话流 (大模型接管)")
async def chat_with_ai(
    req: ChatRequest,
    token: str = Depends(verify_token),
    session: Session = Depends(get_session)
):
    return await process_chat_message(req.record_id, req.reporter_text, session)


@router.post("/chat-audio", response_model=ChatAudioResponse, summary="3. 语音输入并自动进入 AI 对话")
async def chat_with_audio(
    record_id: int = Form(...),
    file: UploadFile = File(..., description="请上传本轮被询问人的录音文件"),
    token: str = Depends(verify_token),
    session: Session = Depends(get_session)
):
    transcript = (await transcribe_upload_file(file)).strip()
    chat_response = await process_chat_message(record_id, transcript, session)

    return ChatAudioResponse(
        transcript=transcript,
        ai_reply=chat_response.ai_reply,
        status=chat_response.status,
        extracted_info=chat_response.extracted_info
    )


def generate_person_intro(person_info: dict, person_type: str = "被询问人") -> str:
    """根据个人信息JSON生成书面化个人介绍（作为被询问人的回答）"""
    parts = []

    name = person_info.get("姓名", "")
    gender = person_info.get("性别", "")
    age = person_info.get("年龄", "")
    birth_date = person_info.get("出生日期", "")
    id_number = person_info.get("身份证件号码", "")
    is_npc = person_info.get("是否为人大代表", "")
    address = person_info.get("现住址", "")
    contact = person_info.get("联系方式", "")
    household = person_info.get("户籍所在地", "")

    # 姓名
    if name:
        parts.append(f"我叫{name}")

    # 性别和年龄
    gender_age = []
    if gender:
        gender_age.append(f"性别{gender}")
    if age:
        gender_age.append(f"今年{age}岁")
    if gender_age:
        parts.append("，".join(gender_age))

    # 出生日期
    if birth_date:
        parts.append(f"出生于{birth_date}")

    # 身份证号
    if id_number:
        parts.append(f"身份证号码{id_number}")

    # 是否人大代表
    if is_npc:
        if is_npc in ["是", "是的", "人大代表", "□是☑ 否"]:
            parts.append("系人大代表")
        elif is_npc in ["否", "不是", "无", "□是 ☑否"]:
            parts.append("非人大代表")

    # 现住址
    if address:
        parts.append(f"现住{address}")

    # 联系方式
    if contact:
        parts.append(f"联系电话{contact}")

    # 户籍所在地
    if household:
        parts.append(f"户籍所在地{household}")

    if not parts:
        return ""

    intro = "，".join(parts) + "。"
    return intro


async def formalize_text_async(text: str) -> str:
    """使用LLM将语音转文字的口语化文本书面化，去除重复和纠正错字"""
    if not text or len(text.strip()) < 2:
        return text

    try:
        response = await client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": """你是一个专业的文字编辑，负责将口语化的语音转文字内容转换为书面化、工整的表达。

要求：
1. 去除口语中的重复、结巴、语气词（如"嗯"、"啊"、"那个"、"就是"、"然后然后"等）；
2. 修正明显的语音识别错字（如"公案"应为"公安"、"行警"应为"刑警"等）；
3. 去除重复表达的句子或段落，保留最完整的一次；
4. 保持原意不变，用更正式书面化的表达；
5. 不要添加任何解释或说明，只输出转换后的文字；
6. 对于数字、日期、地址等关键信息，确保准确无误。"""
                },
                {
                    "role": "user",
                    "content": f"请将以下口语化文本转换为书面化表达，注意去除重复和纠正错字，只输出转换结果：\n{text}"
                }
            ],
            temperature=0.2,
            max_tokens=500
        )
        return response.choices[0].message.content.strip()
    except Exception:
        # 如果LLM调用失败，返回原文
        return text


def formalize_text_sync(text: str) -> str:
    """同步版本的文本书面化（用于非异步场景）"""
    if not text or len(text.strip()) < 2:
        return text

    try:
        import openai
        sync_client = openai.OpenAI(api_key=LLM_API_KEY, base_url=LLM_BASE_URL)
        response = sync_client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": """你是一个专业的文字编辑，负责将口语化的语音转文字内容转换为书面化、工整的表达。

要求：
1. 去除口语中的重复、结巴、语气词（如"嗯"、"啊"、"那个"、"就是"、"然后然后"等）；
2. 修正明显的语音识别错字（如"公案"应为"公安"、"行警"应为"刑警"等）；
3. 去除重复表达的句子或段落，保留最完整的一次；
4. 保持原意不变，用更正式书面化的表达；
5. 不要添加任何解释或说明，只输出转换后的文字；
6. 对于数字、日期、地址等关键信息，确保准确无误。"""
                },
                {
                    "role": "user",
                    "content": f"请将以下口语化文本转换为书面化表达，注意去除重复和纠正错字，只输出转换结果：\n{text}"
                }
            ],
            temperature=0.2,
            max_tokens=500
        )
        return response.choices[0].message.content.strip()
    except Exception:
        return text


def format_transcript(record: Record, create_info: RecordCreateInfo | None) -> dict:
    extracted = parse_extracted_info(record.extracted_info)
    person_type = create_info.person_type.strip() if create_info and create_info.person_type else "被询问人"
    person_labels = ["嫌疑人", "证人", "目击者", "受害人", "被询问人", person_type]

    qa_pairs = []
    current_question = None
    answer_pattern = re.compile(r"^(" + "|".join(re.escape(label) for label in set(person_labels) if label) + r")：")

    for line in record.content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("AI警官："):
            if current_question:
                qa_pairs.append({"question": current_question, "answer": ""})
            current_question = line.replace("AI警官：", "", 1)
        elif answer_pattern.match(line):
            answer = answer_pattern.sub("", line, count=1)
            if current_question:
                qa_pairs.append({"question": current_question, "answer": answer})
                current_question = None
            else:
                qa_pairs.append({"question": "", "answer": answer})
        elif line.startswith("案件基础信息："):
            continue
        elif re.match(r"^(案件类型|案件名称|被询问人身份|被询问人姓名|证件类型|证件号码)：", line):
            continue

    if current_question:
        qa_pairs.append({"question": current_question, "answer": ""})

    header = {
        "title": "询 问 笔 录",
        "case_name": create_info.case_name if create_info else (extracted.get("案情") or record.title or ""),
        "case_type": create_info.case_type if create_info else "",
        "record_time": record.created_at.strftime("%Y年%m月%d日 %H时%M分") if record.created_at else "",
        "record_location": "新城区公安分局刑警大队办公室",
        "interrogator": f"警号 {record.police_number}",
        "recorder": f"警号 {record.police_number}",
    }

    person_info = {
        "姓名": extracted.get("姓名") or (create_info.person_name if create_info else ""),
        "性别": extracted.get("性别", ""),
        "年龄": extracted.get("年龄", ""),
        "出生日期": extracted.get("出生日期", ""),
        "身份证件号码": extracted.get("身份证件号码") or (create_info.id_number if create_info else ""),
        "是否为人大代表": extracted.get("是否为人大代表", ""),
        "现住址": extracted.get("现住址", ""),
        "联系方式": extracted.get("联系方式", ""),
        "户籍所在地": extracted.get("户籍所在地", ""),
    }

    # 生成书面化个人介绍
    person_intro = generate_person_intro(person_info, person_type)

    case_info = {
        "案情": extracted.get("案情", ""),
        "发生时间": extracted.get("发生时间", ""),
        "发生地点": extracted.get("发生地点", ""),
        "案发经过": extracted.get("案发经过", ""),
        "嫌疑人特征": extracted.get("嫌疑人特征", ""),
        "作案工具及涉案物品": extracted.get("作案工具及涉案物品", ""),
        "其他线索": extracted.get("其他线索", ""),
        "相关人员信息": extracted.get("相关人员信息", ""),
    }

    return {
        "header": header,
        "person_info": person_info,
        "person_intro": person_intro,
        "case_info": case_info,
        "qa_pairs": qa_pairs,
        "status": record.status,
    }


@router.get("/{record_id}/transcript", summary="4. 获取格式化笔录数据（前端预览用）")
async def get_transcript(
    record_id: int,
    token: str = Depends(verify_token),
    session: Session = Depends(get_session)
):
    db_record = session.get(Record, record_id)
    if not db_record:
        raise HTTPException(status_code=404, detail="找不到该笔录ID")

    statement = select(RecordCreateInfo).where(RecordCreateInfo.record_id == record_id)
    create_info = session.exec(statement).first()

    return {
        "code": 200,
        "message": "获取笔录成功",
        "data": format_transcript(db_record, create_info)
    }


@router.put("/{record_id}/transcript", summary="4.1 更新笔录内容（手动编辑）")
async def update_transcript(
    record_id: int,
    update_data: dict,
    token: str = Depends(verify_token),
    session: Session = Depends(get_session)
):
    """更新笔录的个人信息和案件要素信息"""
    db_record = session.get(Record, record_id)
    if not db_record:
        raise HTTPException(status_code=404, detail="找不到该笔录ID")

    # 获取当前的 extracted_info
    try:
        extracted_info = json.loads(db_record.extracted_info) if db_record.extracted_info else {}
    except json.JSONDecodeError:
        extracted_info = {}

    # 更新个人信息字段
    person_info = update_data.get("person_info", {})
    field_mapping = {
        "姓名": "姓名",
        "性别": "性别",
        "年龄": "年龄",
        "出生日期": "出生日期",
        "身份证件号码": "身份证件号码",
        "是否为人大代表": "是否为人大代表",
        "现住址": "现住址",
        "联系方式": "联系方式",
        "户籍所在地": "户籍所在地",
    }
    for ui_field, db_field in field_mapping.items():
        if ui_field in person_info:
            extracted_info[db_field] = person_info[ui_field]

    # 更新案件要素信息
    case_info = update_data.get("case_info", {})
    case_field_mapping = {
        "案情": "案情",
        "发生时间": "发生时间",
        "发生地点": "发生地点",
        "案发经过": "案发经过",
        "嫌疑人特征": "嫌疑人特征",
        "作案工具及涉案物品": "作案工具及涉案物品",
        "其他线索": "其他线索",
        "相关人员信息": "相关人员信息",
        # 故意伤害案专用字段
        "冲突原因": "冲突原因",
        "对方动手方式": "对方动手方式",
        "动手方式": "动手方式",
        "是否还手": "是否还手",
        "对方是否还手": "对方是否还手",
        "双方受伤情况": "双方受伤情况",
        "是否使用武器": "是否使用武器",
        "现场证人": "现场证人",
        "伤情描述": "伤情描述",
        "就医情况": "就医情况",
    }
    for ui_field, db_field in case_field_mapping.items():
        if ui_field in case_info:
            extracted_info[db_field] = case_info[ui_field]

    # 更新问答对（如果有）
    qa_pairs = update_data.get("qa_pairs", [])
    if qa_pairs:
        # 重建 content 中的问答记录
        new_content_lines = []
        for pair in qa_pairs:
            if pair.get("question"):
                new_content_lines.append(f"问：{pair['question']}")
            if pair.get("answer"):
                new_content_lines.append(f"答：{pair['answer']}")

        # 保留原有的案件基础信息部分
        original_content = db_record.content or ""
        header_lines = []
        for line in original_content.split("\n"):
            if line.startswith("案件基础信息：") or line.startswith("案件类型：") or line.startswith("案件名称：") or line.startswith("被询问人") or line.startswith("证件"):
                header_lines.append(line)
            elif not line.startswith("问：") and not line.startswith("答："):
                header_lines.append(line)

        # 组合新内容
        db_record.content = "\n".join(header_lines[:6]) + "\n" + "\n".join(new_content_lines)

    # 保存更新
    db_record.extracted_info = json.dumps(extracted_info, ensure_ascii=False)
    session.add(db_record)
    session.commit()
    session.refresh(db_record)

    # 返回更新后的数据
    statement = select(RecordCreateInfo).where(RecordCreateInfo.record_id == record_id)
    create_info = session.exec(statement).first()

    return {
        "code": 200,
        "message": "笔录更新成功",
        "data": format_transcript(db_record, create_info)
    }


@router.get("/{record_id}/export", summary="5. 导出笔录为Word文档")
async def export_record_to_word(
    record_id: int,
    token: str = Depends(verify_token),
    session: Session = Depends(get_session)
):
    import base64
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="后端缺少 python-docx 依赖，请先安装 requirements.txt") from exc

    def set_chinese_font(run, font_name: str, size: int = 12, bold: bool = False):
        """设置中文字体，同时设置西文和东亚文字字体"""
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = font_name  # 西文字体
        # 关键：设置东亚文字字体（中文）
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def base64_to_docx_image(base64_data: str, max_width_cm: float = 4.0, max_height_cm: float = 1.5):
        """将 Base64 图片数据转换为 docx 可用的图片流"""
        try:
            # 移除可能的 data URL 前缀
            if ',' in base64_data:
                base64_data = base64_data.split(',')[1]
            img_data = base64.b64decode(base64_data)
            img_stream = io.BytesIO(img_data)
            return img_stream, Inches(max_width_cm / 2.54), Inches(max_height_cm / 2.54)
        except Exception as e:
            print(f"签名图片转换失败: {e}")
            return None, None, None

    db_record = session.get(Record, record_id)
    if not db_record:
        raise HTTPException(status_code=404, detail="找不到该笔录ID")

    statement = select(RecordCreateInfo).where(RecordCreateInfo.record_id == record_id)
    create_info = session.exec(statement).first()

    # 获取签名数据
    sig_statement = select(Signature).where(Signature.record_id == record_id)
    signatures = session.exec(sig_statement).all()
    signature_map = {sig.signer_type: sig for sig in signatures}

    transcript_data = format_transcript(db_record, create_info)
    header = transcript_data["header"]
    person_info = transcript_data["person_info"]
    case_info = transcript_data["case_info"]
    qa_pairs = transcript_data["qa_pairs"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("询 问 笔 录")
    set_chinese_font(title_run, "黑体", 22, True)

    doc.add_paragraph()
    for line in [
        f"时    间：{header['record_time']}",
        f"地    点：{header['record_location']}",
        f"询 问 人：{header['interrogator']}",
        f"记 录 人：{header['recorder']}",
        f"案件名称：{header['case_name']}",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_chinese_font(run, "仿宋", 12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("被询问人基本信息：")
    set_chinese_font(run, "黑体", 14, True)

    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    person_table_data = [
        ["姓  名", person_info["姓名"], "性  别", person_info["性别"]],
        ["年  龄", person_info["年龄"], "出生日期", person_info["出生日期"]],
        ["身份证件号码", person_info["身份证件号码"], "是否为人大代表", person_info["是否为人大代表"]],
        ["现 住 址", person_info["现住址"], "联系方式", person_info["联系方式"]],
        ["户籍所在地", person_info["户籍所在地"], "", ""],
    ]
    for row_idx, row_data in enumerate(person_table_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, "仿宋", 11)
    table.cell(4, 1).merge(table.cell(4, 3))

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("案件要素摘要：")
    set_chinese_font(run, "黑体", 14, True)

    for key, value in case_info.items():
        if value and value != "未知":
            p = doc.add_paragraph()
            run = p.add_run(f"【{key}】{value}")
            set_chinese_font(run, "仿宋", 12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("询问过程：")
    set_chinese_font(run, "黑体", 14, True)

    for pair in qa_pairs:
        if pair["question"]:
            p = doc.add_paragraph()
            run = p.add_run(f"问：{pair['question']}")
            set_chinese_font(run, "仿宋", 12, True)
        if pair["answer"]:
            p = doc.add_paragraph()
            run = p.add_run(f"答：{pair['answer']}")
            set_chinese_font(run, "仿宋", 12)

    doc.add_paragraph()
    doc.add_paragraph()

    # 结束语
    p = doc.add_paragraph()
    run = p.add_run("以上笔录我已看过（向我宣读过），和我说的相符。")
    set_chinese_font(run, "仿宋", 12)

    doc.add_paragraph()
    doc.add_paragraph()

    # 签名区域 - 带签名图片（仅被询问人签名）
    signature_types = [
        ("被询问人", "被询问人签名（捺手印）："),
    ]

    for signer_type, label_text in signature_types:
        p = doc.add_paragraph()
        run = p.add_run(label_text)
        set_chinese_font(run, "仿宋", 12)

        # 检查是否有该类型的签名
        if signer_type in signature_map:
            sig = signature_map[signer_type]
            img_stream, _, _ = base64_to_docx_image(sig.signature_data)
            if img_stream:
                run = p.add_run()
                run.add_picture(img_stream, width=Cm(4), height=Cm(1.5))
        else:
            run = p.add_run("________________")
            set_chinese_font(run, "仿宋", 12)

        # 添加签名时间
        if signer_type == "被询问人":
            run = p.add_run(f"    {header['record_time']}")
            set_chinese_font(run, "仿宋", 12)

        doc.add_paragraph()

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    person_name = person_info.get("姓名") or "未知"
    file_name = f"询问笔录_{person_name}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{quote(file_name)}"
        }
    )


@router.post("/{record_id}/signature", summary="6. 提交电子签名")
async def submit_signature(
    record_id: int,
    req: SignatureRequest,
    token: str = Depends(verify_token),
    session: Session = Depends(get_session)
):
    """
    提交电子签名
    前端通过 Canvas 采集手写签名，转为 Base64 图片数据后提交
    """
    db_record = session.get(Record, record_id)
    if not db_record:
        raise HTTPException(status_code=404, detail="找不到该笔录ID")

    # 检查笔录是否已结束
    if db_record.status != STATUS_FINISHED:
        raise HTTPException(status_code=400, detail="笔录尚未结束，无法签名")

    # 检查是否已有相同类型的签名
    statement = select(Signature).where(
        Signature.record_id == record_id,
        Signature.signer_type == req.signer_type
    )
    existing_sig = session.exec(statement).first()
    if existing_sig:
        raise HTTPException(status_code=400, detail=f"{req.signer_type}已签名，无需重复签名")

    # 保存签名
    new_signature = Signature(
        record_id=record_id,
        signer_type=req.signer_type,
        signer_name=req.signer_name,
        signature_data=req.signature_data
    )

    session.add(new_signature)
    session.commit()
    session.refresh(new_signature)

    return {
        "code": 200,
        "message": "签名提交成功",
        "data": {
            "signature_id": new_signature.id,
            "signer_type": new_signature.signer_type,
            "signer_name": new_signature.signer_name,
            "signed_at": new_signature.signed_at.strftime("%Y-%m-%d %H:%M:%S")
        }
    }


@router.get("/{record_id}/signatures", summary="7. 获取笔录的所有签名")
async def get_signatures(
    record_id: int,
    token: str = Depends(verify_token),
    session: Session = Depends(get_session)
):
    """
    获取指定笔录的所有签名记录
    """
    db_record = session.get(Record, record_id)
    if not db_record:
        raise HTTPException(status_code=404, detail="找不到该笔录ID")

    statement = select(Signature).where(Signature.record_id == record_id)
    signatures = session.exec(statement).all()

    signature_list = [
        {
            "id": sig.id,
            "signer_type": sig.signer_type,
            "signer_name": sig.signer_name,
            "signed_at": sig.signed_at.strftime("%Y-%m-%d %H:%M:%S")
        }
        for sig in signatures
    ]

    return {
        "code": 200,
        "message": "获取签名列表成功",
        "data": {
            "signatures": signature_list,
            "total": len(signature_list)
        }
    }


@router.get("/{record_id}/signature/{signer_type}", summary="8. 获取指定类型的签名图片")
async def get_signature_image(
    record_id: int,
    signer_type: str,
    token: str = Depends(verify_token),
    session: Session = Depends(get_session)
):
    """
    获取指定笔录中某类型的签名图片（Base64格式）
    signer_type: 被询问人/询问人/记录人
    """
    db_record = session.get(Record, record_id)
    if not db_record:
        raise HTTPException(status_code=404, detail="找不到该笔录ID")

    statement = select(Signature).where(
        Signature.record_id == record_id,
        Signature.signer_type == signer_type
    )
    signature = session.exec(statement).first()

    if not signature:
        raise HTTPException(status_code=404, detail=f"未找到{signer_type}的签名")

    return {
        "code": 200,
        "message": "获取签名图片成功",
        "data": {
            "signer_type": signature.signer_type,
            "signer_name": signature.signer_name,
            "signature_data": signature.signature_data,
            "signed_at": signature.signed_at.strftime("%Y-%m-%d %H:%M:%S")
        }
    }
